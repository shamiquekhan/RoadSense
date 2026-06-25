#!/usr/bin/env python3
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        "<w:tblBorders %s>"
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideV w:val="none"/>'
        "</w:tblBorders>" % nsdecls("w")
    )
    tblPr.append(borders)


def build_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # Page setup (margins)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Global Style setup
    normal_style = doc.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Calibri"
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Read Markdown
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    in_table = False
    table_rows = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Table parsing
        if line.startswith("|"):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            # Process table rows
            parse_and_add_table(doc, table_rows)
            table_rows = []

        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)  # Dark Blue
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)  # Muted Blue
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)  # Dark Grey
        # Blockquote / Note
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            # Simple bold formatting parser
            text = line[2:]
            parse_bold_runs(p, text)
        # Numbered list
        elif line[0].isdigit() and line[1:3] == ". ":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            text = line
            parse_bold_runs(p, text)
        # Section separator
        elif line == "---":
            doc.add_page_break()
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parse_bold_runs(p, line)

        i += 1

    if table_rows:
        parse_and_add_table(doc, table_rows)

    doc.save(docx_path)
    print(f"Successfully generated docx at: {docx_path}")


def parse_bold_runs(p, text):
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if is_bold:
            run = p.add_run(part)
            run.font.bold = True
        else:
            run = p.add_run(part)
        is_bold = not is_bold


def parse_and_add_table(doc, table_rows):
    # Parse table rows into grid
    grid = []
    for r in table_rows:
        # split by | and remove outer
        parts = [p.strip() for p in r.split("|")[1:-1]]
        # skip separator line
        if all(p.replace("-", "").replace(":", "").strip() == "" for p in parts):
            continue
        grid.append(parts)

    if not grid:
        return

    cols = len(grid[0])
    table = doc.add_table(rows=len(grid), cols=cols)
    set_table_borders(table)
    table.autofit = True

    for row_idx, row_data in enumerate(grid):
        row = table.rows[row_idx]
        is_header = row_idx == 0
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

            run = p.add_run(text)
            run.font.name = "Calibri"
            if is_header:
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_background(cell, "1B365D")  # Dark Blue header
            else:
                run.font.size = Pt(10)
                # Alternating row shading
                if row_idx % 2 == 1:
                    set_cell_background(cell, "F5F7FA")  # Off-white

            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_docx.py <md_path> <docx_path>")
        sys.exit(1)
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
